import docx

file = r"C:\Users\Florian\OneDrive - Universite de Montreal\Thése\remerciement\remerciement.docx"
fileLatex = r"C:\Users\Florian\OneDrive - Universite de Montreal\Thése\remerciement\remerciement2.tex"
doc = docx.Document(file)
# with open(file, errors='ignore') as f:


def get_docx_text(filename: str) -> str:
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return ' \par \n'.join(fullText)


def save_to_tex(full_text: str) -> None:
    with open(fileLatex, 'w', encoding="utf-8") as fltx:
        fltx.write(full_text)
"""         for line in doc.paragraphs:
            fltx.write(line.text)
            fltx.write("\n\n")
 """

#save_to_tex(get_docx_text(file))
for para in doc.paragraphs:
    print(para.text)
